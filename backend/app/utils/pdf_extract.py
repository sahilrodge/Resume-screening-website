"""Document text extraction helpers for resumes."""

from __future__ import annotations

import io
import re

from docx import Document
from pypdf import PdfReader

from app.core.exceptions import AppException

MAX_CHARS = 20_000

ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".rtf"}

ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "application/rtf",
    "text/rtf",
}


def _truncate(text: str) -> str:
    return text.strip()[:MAX_CHARS]


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF. Raises if no readable text is found."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise AppException(
            "Could not read PDF content",
            status_code=400,
            code="pdf_read_error",
            details=str(exc),
        ) from exc

    chunks: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(text.strip())

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise AppException(
            "No extractable text found in PDF (it may be a scanned image)",
            status_code=400,
            code="pdf_no_text",
        )
    return _truncate(combined)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file."""
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:  # noqa: BLE001
        raise AppException(
            "Could not read DOCX content",
            status_code=400,
            code="docx_read_error",
            details=str(exc),
        ) from exc

    chunks = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    combined = "\n\n".join(chunks).strip()
    if not combined:
        raise AppException(
            "No extractable text found in DOCX",
            status_code=400,
            code="docx_no_text",
        )
    return _truncate(combined)


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return _truncate(text)
        except UnicodeDecodeError:
            continue
    raise AppException(
        "Could not decode text resume",
        status_code=400,
        code="txt_read_error",
    )


def extract_text_from_rtf(file_bytes: bytes) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError as exc:  # pragma: no cover
        raise AppException(
            "RTF support is not installed on the server",
            status_code=500,
            code="rtf_support_missing",
        ) from exc

    raw = None
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            raw = file_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if raw is None:
        raise AppException(
            "Could not decode RTF resume",
            status_code=400,
            code="rtf_read_error",
        )
    try:
        text = rtf_to_text(raw)
    except Exception as exc:  # noqa: BLE001
        raise AppException(
            "Could not read RTF content",
            status_code=400,
            code="rtf_read_error",
            details=str(exc),
        ) from exc
    text = (text or "").strip()
    if not text:
        raise AppException(
            "No extractable text found in RTF",
            status_code=400,
            code="rtf_no_text",
        )
    return _truncate(text)


def extract_text_from_doc(file_bytes: bytes) -> str:
    """
    Best-effort extraction for legacy .doc (OLE) files.

    Full Word binary parsing is unreliable without LibreOffice; we extract
    readable Unicode/ASCII runs so AI parsing can still run when possible.
    """
    chunks: list[str] = []
    try:
        import olefile

        bio = io.BytesIO(file_bytes)
        if olefile.isOleFile(bio):
            ole = olefile.OleFileIO(bio)
            for stream_name in ("WordDocument", "1Table", "0Table"):
                if not ole.exists(stream_name):
                    continue
                data = ole.openstream(stream_name).read()
                # Prefer UTF-16LE runs common in Word binaries
                try:
                    decoded = data.decode("utf-16-le", errors="ignore")
                    parts = re.findall(r"[\x20-\x7E\u00A0-\u024F]{4,}", decoded)
                    chunks.extend(parts)
                except Exception:  # noqa: BLE001
                    pass
            ole.close()
    except Exception:  # noqa: BLE001
        pass

    if not chunks:
        # Fallback: printable ASCII runs from raw bytes
        ascii_text = "".join(
            chr(b) if 32 <= b < 127 else " " for b in file_bytes
        )
        chunks = [part for part in re.split(r"\s{2,}", ascii_text) if len(part) >= 4]

    combined = "\n".join(chunks).strip()
    combined = re.sub(r"[ \t]+", " ", combined)
    combined = re.sub(r"\n{3,}", "\n\n", combined)
    if len(combined) < 40:
        raise AppException(
            "Could not extract readable text from .doc file. "
            "Please re-save as PDF or DOCX and upload again.",
            status_code=400,
            code="doc_no_text",
        )
    return _truncate(combined)


def detect_extension(filename: str) -> str:
    name = (filename or "").lower().strip()
    if "." not in name:
        return ""
    return "." + name.rsplit(".", 1)[-1]


def detect_file_type(filename: str, content_type: str) -> str:
    ext = detect_extension(filename)
    ctype = (content_type or "").lower()
    if ext == ".pdf" or ctype == "application/pdf":
        return "application/pdf"
    if ext == ".docx" or "wordprocessingml" in ctype:
        return (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    if ext == ".doc" or ctype == "application/msword":
        return "application/msword"
    if ext == ".rtf" or "rtf" in ctype:
        return "application/rtf"
    if ext == ".txt" or ctype.startswith("text/plain"):
        return "text/plain"
    return ctype or "application/octet-stream"


def extract_resume_text(*, file_bytes: bytes, filename: str, content_type: str) -> str:
    """Route extraction by filename / content type."""
    ext = detect_extension(filename)
    ctype = (content_type or "").lower()

    if ext == ".docx" or "wordprocessingml" in ctype:
        return extract_text_from_docx(file_bytes)
    if ext == ".doc" or ctype == "application/msword":
        return extract_text_from_doc(file_bytes)
    if ext == ".txt" or ctype.startswith("text/plain"):
        return extract_text_from_txt(file_bytes)
    if ext == ".rtf" or "rtf" in ctype:
        return extract_text_from_rtf(file_bytes)
    if ext == ".pdf" or ctype == "application/pdf":
        return extract_text_from_pdf(file_bytes)

    raise AppException(
        "Unsupported resume format for text extraction",
        status_code=400,
        code="unsupported_extract_type",
    )
